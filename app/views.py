import re
import PyPDF2
import requests
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response


class PdfView(APIView):

    def read_pdf(self, pdf_url):
        # Download the PDF content
        response = requests.get(pdf_url)
        response.raise_for_status()

        with open('file/pdf/downloaded_resume.pdf', 'wb') as pdf_file:
            pdf_file.write(response.content)

        # Open the downloaded PDF file
        with open('file/pdf/downloaded_resume.pdf', 'rb') as file:
            pdf_reader = PyPDF2.PdfFileReader(file)
            num_pages = pdf_reader.numPages

            data = ""
            for page_number in range(num_pages):
                page = pdf_reader.getPage(page_number)
                text = page.extractText()
                data = data + text

        return data


    def read_doc(self, docx_url):
        response = requests.get(docx_url)
        response.raise_for_status()

        # Save the downloaded DOCX content to a local file
        with open('file/docs/downloaded_document.docx', 'wb') as docx_file:
            docx_file.write(response.content)

        # Read the content of the local DOCX file
        doc = Document('file/docs/downloaded_document.docx')

        content = []
        # Read paragraphs
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)

        # Read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content.append(cell.text)

        return '\n'.join(content)


    def extract_contact_info(self, data):
        # Regular expressions to find name, phone number, and email address
        name_pattern = re.compile(r'\b[A-Z][A-Z\s]+\b')
        phone_pattern = re.compile(r'(?:(?:\+?\d{1,3})?(?:-|\s)?)?\d{10}(?:\s?/\s?\d{10})?')
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        address_patterns = [
            r'\b\d{1,3}\s+[A-Z][a-z]+\s+[A-Za-z]+\b',                   # e.g., 123 Street Name City
            r'Current Address:\s*(.*)',                                 # e.g., Current Address: Some Address
            r'Address:\s*(.*)',                                         # e.g., Address: Some Address
            r'\b[A-Z][A-Za-z\s,-]+\d{5}\b',                             # e.g., Address with postal code
            r'[A-Z]-\d{1,3}\s+[A-Za-z0-9\s,.-]+'                        # e.g., A-123 Street Name City
        ]
        dob_pattern = re.compile(r'\b(?:\d{1,2}[-/]\w{3,9}[-/]\d{2,4}|\d{1,2}/\d{1,2}/\d{2,4})\b')
        education_pattern = re.compile(r'Education\n(.+?)(?=\n\d{4}|$)', re.DOTALL)

        # Find matches
        name_matches = name_pattern.findall(data)
        phone_matches = phone_pattern.findall(data)
        email_matches = email_pattern.findall(data)
        dob_matches = dob_pattern.findall(data)
        education_matches = education_pattern.findall(data)
        addresses_matches = []
        for pattern in address_patterns:
            addresses_matches.extend(re.findall(pattern, data))

        # Assuming the first match found is the primary one
        name = name_matches[0] + name_matches[1] if name_matches else None
        phone = phone_matches[0] if phone_matches else None
        email = email_matches[0] if email_matches else None
        dob = dob_matches[0] if dob_matches else None
        addresses = addresses_matches[0] if addresses_matches else None
        education = education_matches[0] if education_matches else None

        return name, phone, email, addresses, dob, education

    def post(self, request):
        url = request.data.get('url', None)
        file_type = request.data.get('file_type', None)
        if file_type == "pdf":
            data = self.read_pdf(url)
        elif file_type == "docx":
            data = self.read_doc(url)

        try:
            name, phone, email, addresses, dob, education = self.extract_contact_info(data)
        except Exception as e:
            print('Exception is ::>>>', e)
        response_data = {
            "name": name,
            "phone": phone,
            "email": email,
            "addresses": addresses,
            "date_of_birth": dob,
            "education": education,
            "original_content": data
        }
        return Response(response_data)